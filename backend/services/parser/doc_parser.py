# backend/services/parser/doc_parser.py
from typing import List, Tuple, Optional
from docx import Document
from PIL import Image
import easyocr
import io
import gc
import numpy as np
import logging

class DOCXParser:
    """
    A class for parsing DOCX files with text extraction and OCR for embedded images.
    """
    
    def __init__(self, gpu: bool = False):
        """
        Initialize the DOCX parser.
        
        Args:
            gpu: Whether to use GPU for OCR (if available)
        """
        self.gpu = gpu
        self.logger = logging.getLogger(__name__)
        self._easyocr_reader = None
    
    def _get_ocr_reader(self):
        """Lazy loading of EasyOCR reader."""
        if self._easyocr_reader is None:
            try:
                self._easyocr_reader = easyocr.Reader(['en'], gpu=self.gpu)
                self.logger.info("✅ Loaded EasyOCR reader for DOCX parsing")
            except ImportError:
                self.logger.error("❌ EasyOCR not installed. Please install with: pip install easyocr")
                raise
        return self._easyocr_reader
    
    def extract_text(self, docx_path: str) -> str:
        """
        Extracts text from DOCX file and returns as a single string.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            sections = self.extract_text_with_ocr(docx_path)
            return "\n\n".join([text for _, text in sections])
        except Exception as e:
            self.logger.error(f"❌ Failed to extract text from DOCX {docx_path}: {e}")
            raise
    
    def extract_text_with_ocr(self, docx_path: str) -> List[Tuple[int, str]]:
        """
        Extracts text and OCR from a .docx file.
        Returns list of (section_number, text) tuples.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of (section_number, text) tuples
        """
        results = []
        section_num = 1
        
        try:
            doc = Document(docx_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    results.append((section_num, text))
                    section_num += 1
            
            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_text_from_table(table)
                if table_text:
                    results.append((section_num, table_text))
                    section_num += 1
            
            # Extract text from embedded images using OCR
            image_texts = self._extract_text_from_images(doc)
            for text in image_texts:
                if text:
                    results.append((section_num, text))
                    section_num += 1
            
            return results
            
        except Exception as e:
            self.logger.error(f"❌ Failed to extract from DOCX {docx_path}: {e}")
            raise
        finally:
            gc.collect()
    
    def _extract_text_from_table(self, table) -> str:
        """
        Extracts text from a DOCX table.
        
        Args:
            table: DOCX table object
            
        Returns:
            Extracted table text
        """
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    
    def _extract_text_from_images(self, doc: Document) -> List[str]:
        """
        Extracts text from embedded images in DOCX using OCR.
        
        Args:
            doc: DOCX document object
            
        Returns:
            List of extracted text from images
        """
        image_texts = []
        reader = self._get_ocr_reader()
        
        try:
            for rel in doc.part._rels:
                target = doc.part._rels[rel].target_ref
                if target and target.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff")):
                    try:
                        image_blob = doc.part.related_parts[target].blob
                        image_stream = io.BytesIO(image_blob)
                        
                        image = Image.open(image_stream).convert("RGB")
                        image_np = np.array(image)
                        
                        ocr_result = reader.readtext(image_np, detail=0)
                        if ocr_result:
                            image_texts.append("\n".join(ocr_result))
                        
                        # Clean up
                        del image, image_np, image_stream, image_blob
                        
                    except Exception as e:
                        self.logger.warning(f"⚠️ Image OCR failed: {e}")
                        continue
            
            return image_texts
            
        except Exception as e:
            self.logger.error(f"❌ Failed to extract text from images: {e}")
            return []
        finally:
            gc.collect()
    
    def update_parameters(self, gpu: Optional[bool] = None) -> None:
        """
        Update parser parameters.
        
        Args:
            gpu: Whether to use GPU for OCR
        """
        if gpu is not None and gpu != self.gpu:
            self.gpu = gpu
            self._easyocr_reader = None  # Force reload with new GPU setting
            self.logger.info(f"🔄 Updated GPU setting: {gpu}")
            
# from typing import List, Tuple
# from docx import Document
# from PIL import Image
# import easyocr, io, gc
# import numpy as np

# def extract_docx_text_with_ocr(docx_path: str) -> List[Tuple[int, str]]:
#     """
#     Lazily extracts text and OCR from a .docx file.
#     Returns list of (section_number, text) tuples.
#     """
#     doc = Document(docx_path)
#     reader = easyocr.Reader(['en'], gpu=False)
#     results = []
#     section_num = 1

#     # Stream paragraphs
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if text:
#             results.append((section_num, text))
#             section_num += 1

#     # Stream embedded images
#     for rel in doc.part._rels:
#         target = doc.part._rels[rel].target_ref
#         if target.lower().endswith((".png", ".jpg", ".jpeg")):
#             try:
#                 # Lazy load image blob
#                 image_blob = doc.part.related_parts[target].blob
#                 image_stream = io.BytesIO(image_blob)

#                 # Decode only when needed
#                 image = Image.open(image_stream).convert("RGB")
#                 image_np = np.array(image)

#                 ocr_result = reader.readtext(image_np, detail=0)
#                 if ocr_result:
#                     results.append((section_num, "\n".join(ocr_result)))
#                     section_num += 1

#                 del image, image_np, image_stream, image_blob
#                 gc.collect()
#             except Exception as e:
#                 print(f"Image OCR failed: {e}")

#     del reader
#     gc.collect()
#     return results
